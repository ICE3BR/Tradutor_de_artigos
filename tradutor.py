import os

import requests
from docx import Document

subscription_key = "KEY"
endpoint = "https://api.cognitive.microsofttranslator.com"
location = "eastus2"
target_language = "pt-br"


def translator_text(text, target_language):
    path = "/translate"
    constructed_url = endpoint + path
    headers = {
        "Ocp-Apim-Subscription-Key": subscription_key,
        "Ocp-Apim-Subscription-Region": location,
        "Content-type": "application/json",
        "X-ClientTraceId": str(os.urandom(16)),
    }

    body = [{"text": text}]

    params = {"api-version": "3.0", "from": "en", "to": target_language}
    request = requests.post(constructed_url, headers=headers, json=body, params=params)
    response = request.json()
    return response[0]["translations"][0]["text"]


translator_text(
    "I just woke up from a dream, Where you and I had to say goodbye", target_language
)  # Exemplo de uso


def translator_document(path):
    document = Document(path)
    full_text = []
    for paragraph in document.paragraphs:
        translated_text = translator_text(paragraph.text, target_language)
        full_text.append(translated_text)

    translated_doc = Document()
    for line in full_text:
        print(line)
        translated_doc.add_paragraph(line)

    path_translated = path.replace(".docx", f"_{target_language}.docx")
    translated_doc.save(path_translated)
    return path_translated


input_file = "/content/musica.docx"  # Arquivo Docx a ser traduzido
translator_document(input_file)
